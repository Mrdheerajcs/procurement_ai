from difflib import SequenceMatcher
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from docx import Document
from reportlab.pdfgen import canvas

import shutil
import os

from connection import execute_sql_query

from ocr_service import (

    scan_pdf,

    scan_word_document,

    search_documents,

    get_ai_pdf_suggestions,

    get_auto_suggestions
)


DB_URL = 'postgresql://mmudevdb:mmudevdb@103.133.215.182:5432/Procurement'


app = FastAPI()


app.add_middleware(

    CORSMiddleware,

    allow_origins=["*"],

    allow_credentials=True,

    allow_methods=["*"],

    allow_headers=["*"],
)


UPLOAD_FOLDER = "PDFs"

os.makedirs(

    UPLOAD_FOLDER,

    exist_ok=True
)


LAST_GENERATED_TENDER = {}


class TenderRequest(BaseModel):

    requirement: str

    department_id: int | None = None

    mpr_type_id: str | None = None

    document_path: str | None = None


@app.get("/")
def home():

    return FileResponse("index.html")


@app.post("/upload-pdf")

async def upload_pdf(

    file: UploadFile = File(...),

    category: str = "General",

    sub_category: str = "General"
):

    folder_path = os.path.join(

        UPLOAD_FOLDER,

        "Past_Tenders",

        category,

        sub_category
    )

    os.makedirs(

        folder_path,

        exist_ok=True
    )

    file_path = os.path.join(

        folder_path,

        file.filename
    )

    with open(file_path, "wb") as buffer:

        shutil.copyfileobj(

            file.file,

            buffer
        )

    if file.filename.endswith(".pdf"):

        scan_pdf(

            file_path,

            file.filename,

            category,

            sub_category
        )

        return {

            "success": True,

            "message":
            "PDF uploaded and scanned successfully",

            "file_path":
            file_path
        }

    elif file.filename.endswith(".docx"):

        scan_word_document(

            file_path,

            file.filename,

            category,

            sub_category
        )

        return {

            "success": True,

            "message":
            "Word document uploaded and scanned successfully",

            "file_path":
            file_path
        }

    else:

        return {

            "success": False,

            "message":
            "Only PDF and DOCX files are allowed"
        }


@app.get("/search-pdf")
def search_pdf(query: str):

    results = search_documents(query)

    return {

        "matching_files": results
    }


@app.get("/suggestions")
def suggestions(query: str):

    results = get_auto_suggestions(query)

    return {

        "success": True,

        "results": results
    }


@app.post("/api/ai/generate-tender")

def generate_tender(request: TenderRequest):

    global LAST_GENERATED_TENDER

    requirement = request.requirement.lower()

    suggested_items = []

    past_references = []

    total_estimated_value = 0


    suggested_title = (

        request.requirement.title()

        + " Procurement - IT Department"
    )


    sql = """

    SELECT
    d.*,
    h.project_name,
    h.justification,
    h.document_path,
    h.mpr_no
    FROM mpr_details d
    JOIN mpr_header h
    ON d.mpr_id = h.mpr_id
    ORDER BY d.mpr_detail_id DESC
    LIMIT 5
    """


    df = execute_sql_query(

        sql,

        DB_URL
    )


    print("\n========== MPR DETAILS DATA ==========\n")

    print(df.to_string())
    print("\nCOLUMN NAMES:\n")
    print(df.columns.tolist())

    print("\n======================================\n")

    for _, row in df.iterrows():

        quantity = row.get("requested_qty", 1)

        rate = row.get("estimated_rate", 0)

        total_value = quantity * rate

        total_estimated_value += total_value

        suggested_items.append({

            "item_code":
            f"ITM{len(suggested_items)+1:03}",

            "item_name":
            row.get("item_name", "Unknown Item"),

            "quantity":
            quantity,

            "unit":
            "Nos",

            "estimated_rate":
            rate,

            "specification":
            row.get("specificationn", "N/A"),

            "total_value":
            total_value
        })

        search_text = (
            str(row.get("project_name", "")) + " " +
            str(row.get("justification", "")) + " " +
            str(row.get("document_path", ""))
        )

        similarity = int(
            SequenceMatcher(
                None,
                requirement,
                search_text.lower()
            ).ratio() * 100
        )

        past_references.append({
            "mpr_id":row.get("mpr_id", 0),
            "mpr_no":
            row.get(
                "mpr_no",
                f"MPR/2025/00{len(past_references)+1}"
            ),

            "similarity":
            similarity,

            "project_name":
            str(row.get("project_name", ""))
        })
    LAST_GENERATED_TENDER = {
        "success": True,
        "project_name": (
            past_references[0]["project_name"]
            if past_references 
            else ""
        ),
        "suggested_title": suggested_title,
        "suggested_items": suggested_items,
        "total_estimated_value": total_estimated_value,
        "past_references": past_references
    }

    return LAST_GENERATED_TENDER


@app.get("/download-word")
def download_word():

    global LAST_GENERATED_TENDER

    doc = Document()

    doc.add_heading(
        "DETAILED TENDER NOTICE",
        level=1
    )

    doc.add_paragraph(
        "Technical Specification + BOQ + Approval Note"
    )

    doc.add_heading(
        "MATERIAL PURCHASE REQUEST",
        level=2
    )

    if LAST_GENERATED_TENDER.get("past_references"):

        doc.add_paragraph(
            f"MPR No: {LAST_GENERATED_TENDER['past_references'][0]['mpr_no']}"
        )

    doc.add_paragraph(
    f"Project Name: {LAST_GENERATED_TENDER.get('project_name', '')}"
)

    table = doc.add_table(rows=1, cols=7)

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Item No"
    hdr[1].text = "Item Name"
    hdr[2].text = "Mpr Header ID"
    hdr[3].text = "Specification"
    hdr[4].text = "UOM"
    hdr[5].text = "Quantity"
    hdr[6].text = "Estimated Rate"

    for index, item in enumerate(
        LAST_GENERATED_TENDER.get("suggested_items", []),
        start=1
    ):

        row = table.add_row().cells

        row[0].text = str(index)
        row[1].text = str(item.get("item_name", ""))
        row[2].text = str(item.get("mpr_no",""))
        row[3].text = str(item.get("specification", ""))
        row[4].text = str(item.get("unit", "Nos"))
        row[5].text = str(item.get("quantity", 0))
        row[6].text = str(item.get("estimated_rate", 0))

    doc.add_paragraph(
        f"Total Estimated Value: {LAST_GENERATED_TENDER.get('total_estimated_value', 0)}"
    )

    doc.add_heading(
        "TECHNICAL SPECIFICATIONS",
        level=2
    )

    for item in LAST_GENERATED_TENDER.get(
        "suggested_items",
        []
    ):

        doc.add_paragraph(
            f"• {item.get('specification', '')}"
        )

    file_name = "Tender_Report.docx"

    doc.save(file_name)

    return FileResponse(
        file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=file_name
    )


@app.get("/download-pdf")
def download_pdf():

    global LAST_GENERATED_TENDER

    file_name = "Tender_Report.pdf"

    doc = SimpleDocTemplate(
        file_name,
        pagesize=A4
    )

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph(
            "<b>DETAILED TENDER NOTICE</b>",
            styles["Title"]
        )
    )

    elements.append(
        Paragraph(
            "Technical Specification + BOQ + Approval Note",
            styles["BodyText"] 
        )
    )

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            "<b>MATERIAL PURCHASE REQUEST</b>",
            styles["Heading2"]
        )
    )

    if LAST_GENERATED_TENDER.get("past_references"):

        elements.append(
            Paragraph(
                f"MPR No: {LAST_GENERATED_TENDER['past_references'][0]['mpr_no']}",
                styles["BodyText"]
            )
        )

    elements.append(
       Paragraph(
    f"Project Name: {LAST_GENERATED_TENDER.get('project_name', '')}",
    styles["BodyText"]
)
    )

    elements.append(
        Spacer(1, 15)
    )

    table_data = [[
        "Item No",
        "Item Name",
        "Mpr Header ID",
        "Specification",
        "UOM",
        "Quantity",
        "Estimated Rate"
    ]]

    for index, item in enumerate(
        LAST_GENERATED_TENDER.get("suggested_items", []),
        start=1
    ):

        table_data.append([
            str(index),
            str(item.get("item_name", "")),
            str(item.get("mpr_no",0)),
            str(item.get("specification", "")),
            str(item.get("unit", "Nos")),
            str(item.get("quantity", 0)),
            str(item.get("estimated_rate", 0))
            
        ])

    table = Table(table_data)

    table.setStyle(
        TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey)
        ])
    )

    elements.append(table)

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            f"<b>Total Estimated Value:</b> {LAST_GENERATED_TENDER.get('total_estimated_value', 0)}",
            styles["BodyText"]
        )
    )

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            "<b>TECHNICAL SPECIFICATIONS</b>",
            styles["Heading2"]
        )
    )

    for item in LAST_GENERATED_TENDER.get(
        "suggested_items",
        []
    ):

        elements.append(
            Paragraph(
                f"• {item.get('specification', '')}",
                styles["BodyText"]
            )
        )

    doc.build(elements)

    return FileResponse(
        file_name,
        media_type="application/pdf",
        filename=file_name
    )
